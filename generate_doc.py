from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(11)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('知识剧场')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('让学习像追剧一样上瘾')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xF5, 0x57, 0x6C)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('TRAE AI 创造力大赛 · 学习工作赛道')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ==================== 1. 创意名称 + 创意介绍 ====================
h1 = doc.add_heading('1. 创意名称 + 创意介绍', level=1)

doc.add_heading('创意名称', level=2)
p = doc.add_paragraph()
run = p.add_run('知识剧场 — 让学习像追剧一样上瘾')
run.font.size = Pt(14)
run.font.bold = True

doc.add_heading('想解决什么问题', level=2)
doc.add_paragraph(
    '中小学生每天面对千篇一律的练习题和试卷，知识点被拆成孤立的碎片，学生不知道"为什么学"、"学了有什么用"。'
    '长期下来，学习变成了机械重复，热情被消磨殆尽。作业不是太少，而是太无聊。'
)

doc.add_heading('为什么会想到做这个', level=2)
doc.add_paragraph(
    '身边很多中小学生对学习提不起劲，但刷短视频却能刷几个小时。核心区别在于：短视频有故事、有悬念、有情绪起伏，'
    '而作业什么都没有。我在想——如果把知识点编成一个有剧情的故事，让学生在"追剧"的过程中不知不觉把知识学了，会不会有效？'
)

doc.add_heading('大概是什么产品', level=2)
doc.add_paragraph(
    '一个 AI 驱动的 Web 网站（后续可扩展为 App），输入年级和知识点，自动生成"故事化作业"，'
    '像网络短剧一样分集呈现，学生在剧情推进中答题闯关。'
)

# ==================== 2. 目标用户及痛点 ====================
doc.add_heading('2. 目标用户及痛点', level=1)

doc.add_heading('面向哪些用户', level=2)
doc.add_paragraph('小学 1-6 年级学生（核心用户）、他们的家长和老师（决策者/推荐者）')

doc.add_heading('在什么场景下使用', level=2)
doc.add_paragraph('• 放学后完成作业时，用"知识剧场"替代部分传统习题', style='List Bullet')
doc.add_paragraph('• 课堂上老师用来引入新知识点，激发兴趣', style='List Bullet')
doc.add_paragraph('• 周末或假期自主学习，以"追剧"方式系统学习一个章节', style='List Bullet')

doc.add_heading('当前痛点', level=2)
doc.add_paragraph('• 学生：作业枯燥、不想做、做了也不知道有什么用，缺乏内在动力', style='List Bullet')
doc.add_paragraph('• 家长：催作业变成亲子冲突的导火索，孩子不主动学', style='List Bullet')
doc.add_paragraph('• 老师：很难让每个学生都对知识点产生兴趣，个性化趣味教学成本太高', style='List Bullet')

# ==================== 3. 价值与意义 ====================
doc.add_heading('3. 价值与意义', level=1)

doc.add_heading('社会价值', level=2)
doc.add_paragraph(
    '用 AI 技术重新定义"作业"的形态，让学习从被动应付变成主动探索。不是消灭作业，而是让作业变得值得做。'
    '对缓解学生厌学情绪、改善亲子关系有实际意义。'
)

doc.add_heading('效率价值', level=2)
doc.add_paragraph(
    'AI 自动生成故事化内容，老师无需自己编排，一个知识点输入即可得到完整的分集剧本式作业，'
    '大幅降低"趣味教学"的制作门槛。'
)

doc.add_heading('可扩展性', level=2)
doc.add_paragraph(
    '当前 MVP 已实现片段式（单知识点体验）和持久式（整章串联长篇故事）两种模式，覆盖数学学科。'
    '后续可扩展到语文、英语、科学等全学科，甚至跨学科知识融合。'
)

# ==================== 4. 核心功能设计 ====================
doc.add_heading('4. 核心功能设计', level=1)

doc.add_heading('4.1 两种学习模式', level=2)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['', '片段式', '持久式']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['适用场景', '碎片时间快速体验', '系统学习整章知识'],
    ['体验时长', '3-5 分钟/知识点', '20-30 分钟/章节'],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph()

doc.add_heading('4.2 故事化作业流程', level=2)
doc.add_paragraph('① 家长/学生选择年级 + 学科 + 题材风格（武侠/科幻/校园/古风）')
doc.add_paragraph('② AI 基于模板骨架 + 实时生成，输出分集故事')
doc.add_paragraph('③ 学生阅读剧情 → 遇到知识考验 → 答题推进剧情')
doc.add_paragraph('④ 答对：进入主线剧情，解锁下一集')
doc.add_paragraph('⑤ 答错：进入支线引导，讲解知识点后返回主线')
doc.add_paragraph('⑥ 每集开头自动播放"前情回顾"，支持断点续传')

doc.add_heading('4.3 AI 技术方案（模板骨架 + AI 填充）', level=2)

doc.add_paragraph('模板层（保证质量）：', style='List Bullet')
doc.add_paragraph('    • 故事结构骨架（几幕、每幕知识点、答题逻辑）', style='List Bullet')
doc.add_paragraph('    • Prompt 模板（约束 AI 生成范围）', style='List Bullet')
doc.add_paragraph('    • 答题数据（题目、选项、正确答案、解析）', style='List Bullet')

doc.add_paragraph('AI 填充层（保证多样性）：', style='List Bullet')
doc.add_paragraph('    • 每幕剧情文字（根据风格生成不同叙述）', style='List Bullet')
doc.add_paragraph('    • 场景描写（武侠风 vs 科幻风的主要差异）', style='List Bullet')
doc.add_paragraph('    • 角色对话（增加代入感）', style='List Bullet')
doc.add_paragraph('    • 分支过渡语 + 前情回顾摘要', style='List Bullet')

# ==================== 5. 差异化定位 ====================
doc.add_heading('5. 差异化定位', level=1)

doc.add_paragraph(
    '市面上的 AI 辅导产品（豆包、学而思 AI、猿辅导等）核心逻辑是"你遇到问题来问我"，本质是答疑工具。'
)
doc.add_paragraph(
    '知识剧场的核心差异是：从"被动解答"变为"主动想做"。'
)
doc.add_paragraph(
    '别人的产品是"遇到不会的来问我"，我的产品是"还没学就想去做作业"。'
)
doc.add_paragraph(
    '通过故事化包装 + 分支剧情 + 个性化风格，让学生产生内在驱动力，'
    '而不是依赖外部监督。'
)

# ==================== 6. 已实现功能 ====================
doc.add_heading('6. 已实现功能（MVP）', level=1)

doc.add_paragraph('• 三步选择器：年级 → 学科 → 学习模式（片段式/持久式）', style='List Bullet')
doc.add_paragraph('• 短剧式学习页：打字机效果 + 场景描述 + 互动答题', style='List Bullet')
doc.add_paragraph('• 分支剧情：答对推进故事，答错走纠正支线', style='List Bullet')
doc.add_paragraph('• 前情回顾：每集开头自动摘要，支持断点续传', style='List Bullet')
doc.add_paragraph('• 知识地图：可视化学习进度，按章节组织知识点', style='List Bullet')
doc.add_paragraph('• 知识卡片：随时查看当前知识点信息', style='List Bullet')
doc.add_paragraph('• 暗色沉浸式 UI：针对学生群体的视觉设计', style='List Bullet')

# ==================== 7. 问题回顾 ====================
doc.add_heading('7. 核心问题与解决方案', level=1)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['#', '问题', '解决方案']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h

problems = [
    ['1', '知识准确性', '初中小学题目简单，AI 生成计算题出错率极低'],
    ['2', '追剧感如何维持', '分支选择 + AI 生成不同风格故事，每个学生路径不同'],
    ['3', '风格选择困难', '家长根据孩子喜好代选，降低小学生决策负担'],
    ['4', 'AI 接入方式', '模板骨架保证质量 + AI 填充保证多样性'],
    ['5', '核心差异化', '从"被动解答"变为"主动想做"，激发内在驱动力'],
    ['6', '断点续传', '每集开头"前情回顾"，2-3 句话摘要上集剧情'],
]
for r, row_data in enumerate(problems, 1):
    for c, val in enumerate(row_data):
        table2.rows[r].cells[c].text = val

# 保存
output_path = r'E:\Studyproject\codebody\story-homework\知识剧场-TRAE大赛创意文档.docx'
doc.save(output_path)
print(f'文档已保存到: {output_path}')
